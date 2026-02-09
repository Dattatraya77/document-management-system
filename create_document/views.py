import base64
from django.core.files.base import ContentFile
from docx.shared import Mm
from docxtpl import InlineImage, DocxTemplate
from .models import *
from timezonefinder import TimezoneFinder
import datetime
import json
import docx
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.files.storage import default_storage
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db import connection
from django.db.models import Q
from django.http import JsonResponse
from django.shortcuts import redirect, get_object_or_404
from django.shortcuts import render
from django.views import View, generic
from django.views.generic.list import ListView
from django.contrib.auth.models import Group, User
from document_management_system.settings import MEDIA_ROOT
from users.views import get_current_tenant
from users.models import Profile
from create_document.actions import docx_validation, get_meta_text, new_doc_template
from create_document.models import Template, TemplateMetaData, MetadataKey, METADATA_TYPE, CreatedDocument, MetadataValue, MetadataChoice




class SearchDashboard(LoginRequiredMixin, View):

    def get(self, request):
        user = self.request.user
        log_file_path = MEDIA_ROOT + '/' + connection.tenant.name

        # Create the tenant directories
        if connection.tenant.first_visit:
            if not os.path.isdir(log_file_path):
                os.mkdir(log_file_path)
                # Add tenant time zone to the user time zone
                user_profile = get_object_or_404(Profile, user=self.request.user)
                tenant_tz = connection.tenant.client_tz
                user_profile.client_tz = tenant_tz
                user_profile.save()
            if not os.path.isdir(log_file_path + '/created_documents'):
                os.mkdir(log_file_path + '/created_documents')
            if not os.path.isdir(log_file_path + '/templates'):
                os.mkdir(log_file_path + '/templates')
            try:
                Group.objects.create(name='COMMON')
                Group.objects.create(name='DOC_ADMIN')
            except:
                pass

            connection.tenant.first_visit = False
            connection.tenant.save()
        admin_group = get_object_or_404(Group, name='DOC_ADMIN')

        if admin_group in request.user.groups.all():
            template_doc = Template.objects.filter(temp_status='ac').distinct().count()
            created_doc = CreatedDocument.objects.filter(status='ac').exclude(document__exact='').count()
            users = User.objects.all().order_by("id")
            total_user = User.objects.filter(is_active=True).count()
        else:
            common_group = get_object_or_404(Group, name='COMMON')
            template_doc = Template.objects.filter(Q(temp_status='ac',
                                                     doc_group__in=request.user.groups.all()) |
                                                   Q(temp_status='ac',
                                                     doc_group__in=[common_group]) |
                                                   Q(temp_status='ac',
                                                     temp_owner=request.user)).distinct().count()
            created_doc = CreatedDocument.objects.filter(Q(status='ac',
                                                       doc_group__in=request.user.groups.all()) |
                                                     Q(status='ac',
                                                       doc_group__in=[common_group]) |
                                                     Q(status='ac',
                                                       doc_created_by=request.user)).distinct().exclude(document__exact='').count()
            if request.user.is_staff == True:
                users = User.objects.all().order_by("id")
                total_user = User.objects.filter(is_active=True).count()
            else:
                users = User.objects.filter(groups__in=request.user.groups.all()).order_by("id")
                total_user = User.objects.filter(Q(is_active=True, groups__in=request.user.groups.all()) |
                                                 Q(groups__in=[common_group])).distinct().count()

        context = {
            "template_doc": template_doc,
            "created_doc": created_doc,
            "users": users,
            "total_user": total_user,
            'client_name': connection.tenant.name,
            "get_request": True,
        }
        return render(request, 'search_dashboard.html', context)


class GetTimeZone(LoginRequiredMixin, View):
    def get(self, request, *args, **kwargs):
        try:
            latitude = request.GET.get('latitude')
            longitude = request.GET.get('longitude')
            tf = TimezoneFinder()
            tz = tf.timezone_at(lng=float(longitude), lat=float(latitude))
            return JsonResponse({'success': True, 'tz': str(tz)})
        except Exception as e:
            print("error:->",e)
            return JsonResponse({'success': False, 'error': str(e)})


def file_validation(file):
    name = file.name
    if name.endswith('.docx') or name.endswith('.pdf') or name.endswith('.PDF'):
        return True
    else:
        return False


class Index(LoginRequiredMixin, View):

    def post(self, request):
        """
            Step 1 : Find if a matching template exists
            Step 2: Send JsonResponse for matching template alert
            Step 3: Save the template in a media temporary directory,
                    add all the details in session
        """

        response = {
            'success': False,
            'templateID': '',
            'similarStructure': None,
            'debug': '',
        }

        try:
            file_list = request.FILES.getlist('files')
            if len(file_list):
                template = file_list[0]

                from create_document.actions import docx_validation
                if not (docx_validation(template)):
                    response['debug'] = 'Supported file format: Docx'
                    return JsonResponse(response, safe=False)
                # Copy the file
                path = connection.tenant.name + '/templates/' + template.name
                file_name = default_storage.save(path, template)
                request.session['file'] = file_name
                template_obj = docx.Document(template)
                template_para_obj = template_obj.paragraphs[0].text
                response['similarStructure'] = ''
                # Get the timestamp for id
                _now = datetime.datetime.now()
                temporary_id = _now.strftime('%Y%m%d%H%M%S%f')
                request.session['temporary_id'] = temporary_id
                response['template_id'] = temporary_id
                response['success'] = True
            else:
                response['debug'] = 'No file detected'
        except Exception as e:
            response['debug'] = str(e)

        return JsonResponse(response, safe=False)

    def get(self, request):
        try:
            if not request.session['message_checked']:
                if request.session['message']:
                    messages.success(request, request.session['message'])
        except KeyError:
            print("KeyError")
        request.session['message_checked'] = True
        user = self.request.user
        common_group = get_object_or_404(Group, name='COMMON')
        admin_group = get_object_or_404(Group, name='DOC_ADMIN')
        if admin_group in user.groups.all():
            template_doc_list = Template.objects.filter(Q(template_flag='docx', temp_status='ac')|Q(template_flag='docx', temp_status='pe')).order_by('-temp_edited_on')

        else:
            template_doc_list = Template.objects.filter(Q(doc_group__in=request.user.groups.all()) |
                                Q(doc_group__in=[common_group]) |
                                Q(temp_owner=request.user)) \
                                .filter(Q(template_flag='docx', temp_status='ac')|Q(template_flag='docx', temp_status='pe')) \
                                .order_by('-temp_edited_on').distinct()
        # template_doc_list = Template.objects.filter(~Q(temp_status='de')).order_by('-temp_created_at')
        page = request.GET.get('page', 1)
        paginator = Paginator(template_doc_list, 100)
        try:
            template_list = paginator.page(page)
        except PageNotAnInteger:
            template_list = paginator.page(1)
        except EmptyPage:
            template_list = paginator.page(paginator.num_pages)

        user_groups = Group.objects.all()
        # Get time zone
        time_zone = get_object_or_404(Profile, user=request.user).client_tz
        context = {
            'type': 'home',
            'template_doc_list': template_list,
            'user_groups': user_groups,
            'range_page': len(template_list.paginator.page_range),
            'total_number': len(template_doc_list),
            "time_zone": time_zone
        }
        context.update(get_current_tenant(request))
        return render(request, 'index1.html', context)


class AddTemplateDetails(LoginRequiredMixin, View):
    """
        Step 1: Gather the data coming
        Step 2: Add the template Kit_info
    """

    def post(self, request, *args, **kwargs):
        template = request.session['file']
        temp_id = request.session['temporary_id']
        temp_description = request.POST['description']
        temp_title = request.POST['title']
        temp_owner = request.user

        template = Template.objects.create(template=template, temp_id=temp_id,
                                           temp_description=temp_description,
                                           temp_title=temp_title, temp_owner=temp_owner,
                                           )
        groups = request.POST.getlist('group_select')
        for group in groups:
            gp = get_object_or_404(Group, name=group)
            template.doc_group.add(gp)
        template.save()
        return redirect('trainingMetaData', template.temp_id)


class TrainingMetaData(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        response = {
            "success": False,
            "debug": ""
        }

        try:
            template = get_object_or_404(Template, pk=kwargs["template_id"])

            # Data from UI
            provision_data = json.loads(request.POST['provision_data'])
            # Data extracted from DOCX
            meta_session = request.session['metadata']
            # Cache existing metadata keys (fast lookup)
            existing_keys = {
                m.metadata_key: m
                for m in MetadataKey.objects.all()
            }

            for item in meta_session:
                meta_key_name = item["text"].strip()
                meta_id = item["id"]  # ✅ IMPORTANT

                # ------------------------------
                # 1️⃣ CREATE / GET MetadataKey
                # ------------------------------
                if meta_key_name in existing_keys:
                    metadata = existing_keys[meta_key_name]
                else:
                    # ✅ Match by ID, not text
                    ui_meta = next(
                        (m for m in provision_data if m["id"] == meta_id),
                        None
                    )

                    if not ui_meta:
                        print(f"UI metadata not found for id={meta_id}")
                        continue

                    metadata = MetadataKey.objects.create(
                        metadata_key=meta_key_name,
                        metadata_type=ui_meta["type"],
                        metadata_description=ui_meta.get("description", "").strip(),
                        external_metadata=False
                    )

                    existing_keys[meta_key_name] = metadata

                # ------------------------------
                # 2️⃣ CREATE TemplateMetaData
                # ------------------------------
                TemplateMetaData.objects.get_or_create(
                    template=template,
                    metadata_key=metadata,
                    defaults={
                        "temp_description": metadata.metadata_description
                    }
                )

            response["success"] = True
            template.temp_status = "ac"
            template.save()

            request.session["message"] = (
                f"Training template {template.temp_id} completed successfully!"
            )
            request.session["message_checked"] = False

        except Exception as e:
            response["debug"] = str(e)

        return JsonResponse(response, safe=False)

    def get(self, request, *args, **kwargs):
        template = get_object_or_404(Template, pk=kwargs['template_id'])
        document_txt = get_meta_text(template)
        request.session['metadata'] = document_txt['metadata']
        metadata_types = METADATA_TYPE
        existing_meta_keys = MetadataKey.objects.all()
        for key in existing_meta_keys:
            key.metadata_description = key.metadata_description.strip()
        try:
            first_span_id = document_txt['metadata'][0]['id']
            last_span_id = document_txt['metadata'][-1]['id']
        except:
            return redirect('home')

        context = {
            'document_txt': document_txt['text'],
            'template': template,
            "first_span_id": first_span_id,
            "last_span_id": last_span_id,
            "meta_types": metadata_types,
            "existing_meta_keys": existing_meta_keys,
        }
        context.update(get_current_tenant(request))
        return render(request, 'trainingMeta.html', context)


class DownloadViewDocument(LoginRequiredMixin, View):

    def post(self, request):

        document_id = request.POST.get('data')
        document_status = request.POST.get('status')
        document_name = request.POST.get('doc_name')

        if not document_id or not document_status:
            return JsonResponse({'success': False, 'message': 'Invalid request'})

        extension = document_name.split('.')[-1].lower()

        try:
            document_obj = get_object_or_404(CreatedDocument, doc_id=document_id)
            file_field = document_obj.document
        except:
            document_obj = get_object_or_404(Template, temp_id=document_id)
            file_field = document_obj.template

        file_path = file_field.path
        file_url = settings.MEDIA_URL + str(file_field)
        file_size = os.path.getsize(file_path)

        return JsonResponse({
            'success': True,
            'message': 'Document ready',
            'extension': extension,
            'file_url': file_url,
            'file_size': file_size
        })


class DeleteTemplate(LoginRequiredMixin, View):
    """Change the template status to delete"""

    def get(self, request, *args, **kwargs):
        template = get_object_or_404(Template, temp_id=kwargs["template_id"])
        template.temp_status = 'de'
        template.save()
        messages.success(request, "Template Name: {0}".format(template.get_temp_name()) + " deleted successfully.")

        return redirect('train-docx-template')


class TemplatesToDoc(ListView, LoginRequiredMixin):
    model = Template
    template_name = 'templatesToDoc.html'

    def get_queryset(self):
        user = self.request.user
        admin_group = get_object_or_404(Group, name='DOC_ADMIN')
        common_group = get_object_or_404(Group, name='COMMON')
        if admin_group in user.groups.all():
            return Template.objects.filter(temp_status='ac', template_flag='docx').order_by('-temp_edited_on').distinct()
        return Template.objects.filter(Q(temp_status='ac',
                                         doc_group__in=user.groups.all()) |
                                       Q(temp_status='ac',
                                         doc_group__in=[common_group]) |
                                       Q(temp_status='ac', temp_owner=user)).filter(template_flag='docx')\
                               .order_by('-temp_edited_on')\
                               .distinct()

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['client_name'] = connection.tenant.name
        # Get time zone
        time_zone = get_object_or_404(Profile, user=self.request.user).client_tz
        context['time_zone'] = time_zone
        return context


class CreateNewDocumentTemplate(LoginRequiredMixin, View):
    template_name = "createTemplateDocument.html"

    def get(self, request, template_id):
        template = get_object_or_404(Template, pk=template_id)

        ui_text = new_doc_template(template)

        document_name = f"{template.temp_title}-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
        user_groups = Group.objects.all()

        context = {
            "template": template,
            "ui_text": ui_text,
            "document_name": document_name,
            "user_groups": user_groups,
        }
        return render(request, self.template_name, context)

    def post(self, request, template_id):
        template = get_object_or_404(Template, pk=template_id)

        # 1️⃣ Create document record
        doc_id = datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')
        file_name = request.POST.get("file_name") + ".docx"

        created_doc = CreatedDocument.objects.create(
            doc_id=doc_id,
            document=f"{connection.tenant.name}/created_documents/{file_name}",
            doc_matched_template=template,
            doc_created_by=request.user,
            doc_type="docx"
        )

        # Add the document gropus
        groups = request.POST.getlist('group_select')
        for group in groups:
            gp = get_object_or_404(Group, name=group)
            created_doc.doc_group.add(gp)
        created_doc.save()

        # 2️⃣ Prepare DOCX rendering
        tpl = DocxTemplate(template.template.path)
        context = {}

        metadata_keys = TemplateMetaData.objects.filter(template=template)

        for tm in metadata_keys:
            key = tm.metadata_key
            raw_value = request.POST.get(key.metadata_key)

            mv = MetadataValue(
                meta_key=key,
                meta_created_doc=created_doc,
                updated_by=request.user
            )

            if key.metadata_type == "string" or key.metadata_type == "textarea" or key.metadata_type == "choice":
                mv.metadata_value_text = raw_value
                context[key.metadata_key] = raw_value

            elif key.metadata_type == "number":
                mv.metadata_value_number = raw_value or None
                context[key.metadata_key] = raw_value

            elif key.metadata_type == "date":
                mv.metadata_value_date = raw_value or None
                context[key.metadata_key] = raw_value

            elif key.metadata_type == "boolean":
                mv.metadata_value_bool = True if raw_value == "true" else False
                context[key.metadata_key] = raw_value

            elif key.metadata_type == "image":
                file = request.FILES.get(key.metadata_key)
                metadata_key_name = key.metadata_key
                width = metadata_key_name.split('_')[-1].replace("mm", '')
                if file:
                    mv.metadata_value_image = file
                    context[key.metadata_key] = InlineImage(tpl, file, width=Mm(float(width)))

            elif key.metadata_type == "signature":

                # 1️⃣ DRAW signature (canvas → base64)
                value = request.POST.get(key.metadata_key)

                if value and value.startswith("data:image"):
                    format, imgstr = value.split(";base64,")
                    data = ContentFile(
                        base64.b64decode(imgstr),
                        name=f"{key.metadata_key}.png"
                    )

                    mv.metadata_value_image = data
                    context[key.metadata_key] = InlineImage(tpl, data, width=Mm(30))

                # 2️⃣ IMAGE upload signature (file input)
                elif key.metadata_key in request.FILES:
                    file = request.FILES.get(key.metadata_key)
                    mv.metadata_value_image = file
                    context[key.metadata_key] = InlineImage(tpl, file, width=Mm(30))

                # 3️⃣ TYPE signature (text)
                elif value:
                    mv.metadata_value_text = value
                    context[key.metadata_key] = value

            mv.save()

        # 3️⃣ Render and save DOCX
        output_path = os.path.join(
            settings.MEDIA_ROOT,
            connection.tenant.name,
            "created_documents",
            file_name
        )

        tpl.render(context)
        tpl.save(output_path)

        messages.success(request, "Document created successfully")
        return redirect("new-document-list")


class NewDocumentList(LoginRequiredMixin, generic.ListView):
    template_name = 'newDocuments.html'
    context_object_name = 'new_document_list'

    def get_queryset(self):
        user = self.request.user
        common_group = get_object_or_404(Group, name='COMMON')
        admin_group = get_object_or_404(Group, name='DOC_ADMIN')
        if admin_group in user.groups.all():
            documents = CreatedDocument.objects.filter(status='ac', doc_type='docx') \
                .order_by('-doc_updated_on').distinct()
            return documents

        documents = CreatedDocument.objects.filter(Q(doc_group__in=user.groups.all()) |
                                               Q(doc_group__in=[common_group]) |
                                               Q(doc_created_by=user)) \
            .filter(status='ac', doc_type='docx') \
            .order_by('-doc_updated_on').distinct()
        return documents

    def get_context_data(self, *, object_list=None, **kwargs):
        user = self.request.user
        context = super().get_context_data(**kwargs)
        context['client_name'] = connection.tenant.name
        time_zone = get_object_or_404(Profile, user=self.request.user).client_tz
        context['time_zone'] = time_zone
        context['users'] = User.objects.all()
        context['user'] = user
        context['time_zone'] = time_zone
        return context


class DeleteDocxDocument(LoginRequiredMixin, View):
    def get(self, request, *args, **kwargs):
        document = get_object_or_404(CreatedDocument, doc_id=kwargs["doc_id"])
        document.status = 'de'
        document.save()
        messages.success(request, "Created Document ID(" + document.doc_id + ") deleted successfully.")
        return redirect('new-document-list')