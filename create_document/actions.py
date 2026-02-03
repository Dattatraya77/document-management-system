import re
from docx2python import docx2python
from document_management_system.settings import MEDIA_ROOT
from docx import Document
from .models import *
from docx.oxml.ns import qn


def docx_validation(file):
    name = file.name
    if name.endswith('.docx'):
        return True
    else:
        return False


def get_meta_text(template_document, update=[], ui=False):
    path = MEDIA_ROOT + '/' + str(template_document.template)
    document = docx2python(path, html=True)
    text = document.text
    header_text = []
    header = document.header
    for h in header:
        header_text.append(h[0][0][0])

    for h_text in header_text:
        text = re.sub(h_text, '', text)

    text = re.sub(r'(<font .*?>|</font>)', '', text)
    text = re.sub(r'\t', '', text)
    if not ui:
        text = re.sub('</b>', '', text)
        text = re.sub('<b>', '', text)
        text = re.sub(r'<u>', '', text)
        text = re.sub(r'</u>', '', text)
    text_list1 = list(text.split('\n\n'))
    text_list2 = list(filter(''.__ne__, text_list1))
    text_list = list(filter(('\t').__ne__, text_list2))

    metadata = []
    para_text = ''
    pre_text_list1 = []
    post_text_list = []
    for i, para in enumerate(text_list):
        if para.find("{{", 0) != -1:
            key = "para" + str(i)
            pointer = 0
            counter = 0
            while True:
                # find the start, end location for the metdata key
                start_location = para.find("{{", pointer)
                end_location = para.find("}}", start_location + 2)
                if start_location == -1 or end_location == -1:
                    break
                else:
                    text = para[start_location + 2:end_location]
                    text = text.replace(" ", "")
                    pre_text = para[pointer:start_location]
                    # Get the pre text
                    pre_text_list = pre_text.split()
                    if len(pre_text_list) >= 2:
                        pre_text_meta = " ".join(pre_text_list[-6:])
                    elif len(pre_text_list) == 0:
                        if i == 0 or start_location != 0:
                            pre_text_meta = ""
                        else:
                            pre_text_sub = " ".join(text_list[i - 1].split()[-2:])
                            pre_text_meta = re.sub("{{.*?}}", "", pre_text_sub)
                            pre_text_meta = re.sub("}} ", "", pre_text_meta)
                    else:
                        pre_text_meta = " ".join(pre_text_list)

                    j = 0
                    while '}}' in pre_text_meta or '{{' in pre_text_meta:
                        j += 1
                        try:
                            pre_text_meta = ' '.join(text_list[i - j].split()[-2:])
                        except:
                            pass
                    post_text_index = para.find("{{", end_location + 2)

                    if post_text_index != -1:
                        pos_text = para[end_location + 2: post_text_index]
                        pos_text_list = pos_text.split()
                    else:
                        pos_text = para[end_location + 2:]
                        pos_text_list = pos_text.split()

                    if len(pos_text_list) == 0:
                        if i == len(text_list) - 1 or post_text_index == end_location + 3:
                            pos_text_meta = ""
                        else:
                            try:
                                pos_text_meta = text_list[i + 1].split()[0]
                            except:
                                pos_text_meta = ""
                    elif len(pos_text_list) >= 2:
                        pos_text_meta = " ".join(pos_text_list[0:2])

                    else:
                        pos_text_meta = " ".join(pos_text_list)
                    k = 1
                    while '{{' in pos_text_meta or '}}' in pos_text_meta:
                        k += 1
                        pos_text_meta = ' '.join(text_list[i + k].split()[:1])

                    metadata.append({'text': text, 'pre_text': pre_text_meta, 'post_text': pos_text_meta,
                                     'id': str(i) + str(counter)
                                     })

                    pre_text_list1.append(pre_text_meta)
                    post_text_list.append(pos_text_meta)
                    if not ui:
                        if counter == 0:
                            para_text += "<p>" + pre_text + "<span id='span" + str(i) + str(counter) + "'> {{" + text + \
                                         "}}" + "</span>" + pos_text
                        else:
                            para_text += "<span id='span" + str(i) + str(counter) + "'> {{" + text + \
                                         "}}" + "</span>" + pos_text

                    elif len(update) > 0:
                        span_text = re.sub('</b>', '', text)
                        span_text = re.sub('<b>', '', span_text)
                        span_text = re.sub(r'<u>', '', span_text)
                        span_text = re.sub(r'</u>', '', span_text)

                        res = ''
                        for sub in update:
                            if sub['key'] == span_text:
                                res = sub['value']
                                if res == 'None':
                                    res = ''
                                break

                        if counter == 0:
                            if sub['type'] == 'image':

                                para_text += "<p>" + pre_text + "<img src='" + res + "'" \
                                                                                     "width='" + sub[
                                                 'width'] + "px' class='meta_init " + span_text + "'" + \
                                             ">" + pos_text
                            else:
                                para_text += "<p>" + pre_text + "<span class='meta_init " + span_text + "'" + ">" + \
                                             res + "</span>" + pos_text
                        else:
                            if sub['type'] == 'image':
                                para_text += "<img src='" + res + "' class='meta_init " + span_text + "'>" + \
                                             pos_text
                            else:
                                para_text += "<span class='meta_init " + span_text + "'>" + \
                                             res + "</span>" + pos_text
                    else:
                        span_text = re.sub('</b>', '', text)
                        span_text = re.sub('<b>', '', span_text)
                        span_text = re.sub(r'<u>', '', span_text)
                        span_text = re.sub(r'</u>', '', span_text)
                        if counter == 0:
                            if span_text.find('image') == -1:
                                para_text += "<p>" + pre_text + "<span class='meta_init " + span_text + "'" + ">" + \
                                             "______" + "</span>" + pos_text
                            else:
                                para_text += "<p>" + pre_text + "<img src='#' class='" + span_text + "'>" + pos_text

                        else:
                            if span_text.find('image') == -1:
                                para_text += "<span class='meta_init " + span_text + "'>" + \
                                             "______" + "</span>" + pos_text
                            else:
                                para_text += "<img src='#' class='" + span_text + "'>" + pos_text
                    pointer = end_location + 2
                    counter += 1
            para_text += "</p>"
        else:
            para_text += "<p>" + para + "</p>"
    if ui:
        return para_text

    return {'text': para_text, 'metadata': metadata}


def new_doc_template(template):
    doc = Document(template.template.path)

    meta_map = {
        m.metadata_key: m
        for m in MetadataKey.objects.all()
    }

    html_blocks = []

    def render_text(text):
        if not text.strip():
            return ""

        def repl(match):
            key_name = match.group(1)
            meta = meta_map.get(key_name)

            if not meta:
                return ""


            placeholder = meta.metadata_description or key_name.replace("_", " ").title()

            if meta.metadata_type == "date":
                input_html = f"""
                <input type="date"
                       name="{key_name}"
                       class="doc-input date-input"
                       placeholder="{placeholder}">
                """
            elif meta.metadata_type == "number":
                input_html = f"""
                <input type="number"
                       name="{key_name}"
                       class="doc-input number-input"
                       placeholder="{placeholder}">
                """
            elif meta.metadata_type == "boolean":
                input_html = f"""
                <select name="{key_name}" class="doc-input select-input">
                    <option value="">— Select —</option>
                    <option value="true">True</option>
                    <option value="false">False</option>
                </select>
                """
            elif meta.metadata_type == "choice":
                options_html = '<option value="">— Select —</option>'

                for choice in meta.metadata_choice.all():
                    options_html += f'''
                        <option value="{choice.meta_choice}">
                            {choice.meta_choice}
                        </option>
                    '''

                input_html = f"""
                <select name="{key_name}"
                        class="doc-input select-input"
                        data-meta-type="choice">
                    {options_html}
                </select>
                """
            elif meta.metadata_type == "textarea":
                input_html = f"""
                <textarea name="{key_name}"
                          class="doc-input textarea-input"
                          rows="1"
                          placeholder="{placeholder}"></textarea>
                """
            elif meta.metadata_type == "image":
                input_html = f"""
                <input type="file"
                       name="{key_name}"
                       class="doc-input file-input"
                       accept="image/*">
                """
            elif meta.metadata_type == "signature":
                input_html = f"""
                <input type="hidden" name="{key_name}" id="{key_name}_value">

                <button type="button"
                        class="signature-btn"
                        data-key="{key_name}">
                    ✍ Add Signature
                </button>

                <div id="{key_name}_preview" class="signature-preview"></div>
                """
            else:
                input_html = f"""
                <input type="text"
                       name="{key_name}"
                       class="doc-input text-input"
                       placeholder="{placeholder}">
                """

            return input_html

        return re.sub(r"\{\{\s*(\w+)\s*\}\}", repl, text)

    # -------------------------
    # Header
    # -------------------------
    header = doc.sections[0].header
    if header:
        html_blocks.append('<div class="doc-header">')
        for p in header.paragraphs:
            html_blocks.append(f"<p>{render_text(p.text)}</p>")
        html_blocks.append("</div>")

    # -------------------------
    # Helpers
    # -------------------------
    def has_page_break(para):
        for run in para.runs:
            for br in run._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True
        return False

    def get_numbering_prefix(para):
        pPr = para._p.pPr
        if pPr is None or pPr.numPr is None:
            return ""

        ilvl = pPr.numPr.ilvl.val
        numId = pPr.numPr.numId.val

        # simple reconstruction
        if ilvl == 0:
            return "• "
        elif ilvl == 1:
            return "a) "
        elif ilvl == 2:
            return "i) "
        return ""

    # -------------------------
    # Body (paragraphs + tables in order)
    # -------------------------
    body = doc.element.body

    for element in body:
        tag = element.tag.split("}")[-1]

        # -------- Paragraph --------
        if tag == "p":
            para = next(p for p in doc.paragraphs if p._element == element)

            if has_page_break(para):
                html_blocks.append('<div class="page-break"></div>')
                continue

            prefix = get_numbering_prefix(para)
            text = prefix + render_text(para.text)

            if text.strip():
                html_blocks.append(f"<p>{text}</p>")
            else:
                html_blocks.append("<br>")

        # -------- Table --------
        elif tag == "tbl":
            table = next(t for t in doc.tables if t._element == element)

            is_signature = (
                    len(table.columns) == 2 and
                    any("By:" in c.text or "Name:" in c.text for r in table.rows for c in r.cells)
            )

            cls = "doc-table signature-table" if is_signature else "doc-table"
            html_blocks.append(f'<table class="{cls}">')

            for row in table.rows:
                html_blocks.append("<tr>")
                for cell in row.cells:
                    cell_html = "<br>".join(
                        render_text(p.text) for p in cell.paragraphs if p.text.strip()
                    )
                    html_blocks.append(f"<td>{cell_html}</td>")
                html_blocks.append("</tr>")

            html_blocks.append("</table>")

    # -------------------------
    # Footer (first section)
    # -------------------------
    footer = doc.sections[0].footer
    if footer:
        html_blocks.append('<div class="doc-footer">')
        for p in footer.paragraphs:
            html_blocks.append(f"<p>{render_text(p.text)}</p>")
        html_blocks.append("</div>")

    return "\n".join(html_blocks)